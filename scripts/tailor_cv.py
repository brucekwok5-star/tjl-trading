#!/usr/bin/env python3
"""
tailor_cv.py — for each job folder, copy the CV template and update the
profile to match the job description.

Output folder: ~/.openclaw/workspace/job-search/JDs/{Source}/{Company}_{Title}_{hash}/
  ├── job.txt           ← job details (already there)
  ├── cv.docx           ← tailored CV (created by this script)
  └── cover_letter.docx ← placeholder (from save_jds.py)

Usage:
    python3 tailor_cv.py                      # process all job folders
    python3 tailor_cv.py --folder "path/to/folder"  # single folder
    python3 tailor_cv.py --list               # show pending folders
"""

import shutil
import json
import urllib.request
import urllib.error
import sys
import re
from pathlib import Path
from datetime import datetime

# ---------------------------------------------------------------------------
# Config
# ---------------------------------------------------------------------------
HOME = Path.home()
WORKSPACE   = HOME / ".openclaw" / "workspace" / "job-search"
TEMPLATE_CV = WORKSPACE / "Bruce Kwok_Resume.docx"
JD_DIR      = WORKSPACE / "JDs"
MINIMAX_KEY = "sk-cp-Ssq7KhTUX8bJJnroMymIFBn87GWi3K3fmfHpJ2poY4nI5MUUFPeVknVRwI9nCl2SqmfU2kQ-rQwRuRUmZDXUWOuZE_Nvl-voI3yTabGu5C-dK-KhCSA1GbA"
MINIMAX_URL = "https://api.minimaxi.com/anthropic/v1/messages"
MODEL       = "MiniMax-M2.7"
MAX_TOKENS  = 1000

# Current CV profile (used as base context)
CURRENT_PROFILE = (
    "A senior technology leader with 15+ years of IT experience, including 10+ years "
    "in IT infrastructure and 5+ years in team management. Proven ability to lead and "
    "mentor technical teams delivering resilient, secure, and highly available services "
    "across AWS, Microsoft Azure, and on-premises platforms. Deep expertise in cloud "
    "governance, infrastructure automation using CI/CD and IaC (Terraform, CDK, Jenkins, "
    "GitHub Enterprise), and unified management across multi-cloud environments. Skilled "
    "in cloud audit, risk assessment, and compliance."
)

# ---------------------------------------------------------------------------
# MiniMax API
# ---------------------------------------------------------------------------
def generate_profile(job_title: str, job_desc: str) -> str:
    """Call MiniMax to rewrite the profile for the given job."""
    prompt = f"""Rewrite this CV profile for the job. Output ONLY the profile text: 3 sentences, first person, no quotes or labels.
Original: {CURRENT_PROFILE[:200]}
Job: {job_title}
Keywords: {job_desc[:500]}
Profile:"""

    body = json.dumps({
        "model": MODEL,
        "max_tokens": MAX_TOKENS,
        "messages": [{"role": "user", "content": prompt}]
    }).encode()

    req = urllib.request.Request(
        MINIMAX_URL,
        data=body,
        headers={
            "Authorization": f"Bearer {MINIMAX_KEY}",
            "Content-Type": "application/json",
            "x-api-key": MINIMAX_KEY,
            "anthropic-version": "2023-06-01",
        },
        method="POST"
    )

    try:
        with urllib.request.urlopen(req, timeout=60) as resp:
            data = json.loads(resp.read())
    except urllib.error.HTTPError as e:
        body = e.read().decode() if e.fp else ""
        raise RuntimeError(f"MiniMax HTTP {e.code}: {body[:200]}") from e

    # Extract text from content blocks — skip thinking blocks
    text = ""
    for block in data.get("content", []):
        if block.get("type") == "text":
            candidate = block["text"].strip()
            if candidate:
                text = candidate

    if not text:
        raise RuntimeError(f"No text in MiniMax response (stop={data.get('stop_reason')})")
    return text


# ---------------------------------------------------------------------------
# Docx manipulation
# ---------------------------------------------------------------------------
def update_cv(cv_path: Path, job_title: str, new_profile: str):
    """Open cv_path, update title and profile paragraphs, save."""
    from docx import Document

    doc = Document(str(cv_path))

    # Paragraph 2 = job title, Paragraph 3 = profile
    p_title   = doc.paragraphs[2]
    p_profile = doc.paragraphs[3]

    # Update title (single run)
    if p_title.runs:
        p_title.runs[0].text = job_title

    # Replace profile: clear extra runs, set first run to new text
    if p_profile.runs:
        for run in p_profile.runs[1:]:
            run.text = ""
        p_profile.runs[0].text = new_profile

    doc.save(str(cv_path))


# ---------------------------------------------------------------------------
# Per-job processing
# ---------------------------------------------------------------------------
def process_folder(job_dir: Path) -> dict:
    """Tailor CV for one job folder. Returns dict with results."""
    job_txt = job_dir / "job.txt"
    if not job_txt.exists():
        return {"folder": job_dir.name, "status": "skip", "reason": "no job.txt"}

    # Read job details
    lines = job_txt.read_text(encoding="utf-8").splitlines()
    job_info = {}
    for line in lines:
        if ": " in line:
            key, _, val = line.partition(": ")
            job_info[key.strip()] = val.strip()

    job_title   = job_info.get("Title", "Professional")
    job_link    = job_info.get("Link", "")
    source      = job_info.get("Source", "")
    company     = job_info.get("Company", "")

    # Extract JD text (everything after the separator)
    try:
        sep_idx = lines.index("=" * 60)
        jd_text = " ".join(lines[sep_idx + 1:])
        # Clean HTML artifacts
        jd_text = re.sub(r'\[View job posting\]\([^)]+\)', '', jd_text)
        jd_text = re.sub(r'## Notes.*', '', jd_text, flags=re.DOTALL)
        jd_text = re.sub(r'## Your Application.*', '', jd_text, flags=re.DOTALL)
        jd_text = jd_text.strip()
    except ValueError:
        jd_text = job_title + " " + job_info.get("Company", "")

    # Truncate JD text to avoid huge prompts
    jd_text = jd_text[:1200]

    # Destination CV path
    cv_path = job_dir / "cv.docx"

    # Copy template over
    shutil.copy2(TEMPLATE_CV, cv_path)

    # Generate tailored profile
    try:
        new_profile = generate_profile(job_title, jd_text)
    except Exception as e:
        # If API fails, keep template CV but note the error
        return {
            "folder": job_dir.name,
            "status": "error",
            "reason": str(e),
            "cv_path": str(cv_path),
            "job_title": job_title,
            "company": company,
        }

    # Update the CV
    try:
        update_cv(cv_path, job_title, new_profile)
    except Exception as e:
        return {
            "folder": job_dir.name,
            "status": "error",
            "reason": f"CV update failed: {e}",
            "cv_path": str(cv_path),
        }

    return {
        "folder": job_dir.name,
        "status": "ok",
        "job_title": job_title,
        "company": company,
        "cv_path": str(cv_path),
        "profile": new_profile[:100] + "...",
    }


def find_job_folders() -> list[Path]:
    """Find all job folders (depth=2: Source/Company_Title_hash/)."""
    folders = []
    if not JD_DIR.exists():
        return folders
    for source_dir in JD_DIR.iterdir():
        if source_dir.is_dir():
            for job_dir in source_dir.iterdir():
                if job_dir.is_dir():
                    folders.append(job_dir)
    return sorted(folders)


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------
def main():
    import argparse
    parser = argparse.ArgumentParser(description="Tailor CV for each job folder")
    parser.add_argument("--folder", "-f", type=Path, help="Process single folder")
    parser.add_argument("--list", "-l", action="store_true", help="List pending folders")
    parser.add_argument("--force", action="store_true", help="Re-generate even if cv.docx exists")
    args = parser.parse_args()

    if not TEMPLATE_CV.exists():
        print(f"[ERROR] Template CV not found: {TEMPLATE_CV}")
        sys.exit(1)

    if args.list:
        folders = find_job_folders()
        print(f"Job folders ({len(folders)}):")
        for fd in folders:
            cv = fd / "cv.docx"
            status = "✓ cv.docx" if cv.exists() else "○ pending"
            print(f"  {status}  {fd.parent.name}/{fd.name[:50]}")
        return

    if args.folder:
        folders = [args.folder]
    else:
        folders = find_job_folders()

    if not folders:
        print("No job folders found.")
        return

    done = 0
    errors = 0

    for job_dir in folders:
        cv_path = job_dir / "cv.docx"
        if cv_path.exists() and not args.force:
            print(f"  [SKIP] {job_dir.name[:50]}  (cv.docx already exists)")
            continue

        print(f"\nProcessing: {job_dir.parent.name}/{job_dir.name[:50]}")
        result = process_folder(job_dir)

        if result["status"] == "ok":
            print(f"  ✓ {result['job_title']} @ {result['company']}")
            print(f"    Profile: {result['profile']}")
            done += 1
        else:
            print(f"  ✗ {result.get('reason', 'unknown error')}")
            errors += 1

    print(f"\n{'='*60}")
    print(f"Done: {done} CVs tailored, {errors} errors")
    if done:
        print(f"CVs saved in: {JD_DIR}")


if __name__ == "__main__":
    main()
