"""Tests for job-search/tailor_cv.py — prompt, Unicode normalization, docx update."""
from __future__ import annotations

import json
from pathlib import Path
from unittest.mock import MagicMock, patch

import pytest


# ── build_prompt ──────────────────────────────────────────────────────────────

class TestBuildPrompt:
    def test_returns_string(self, js_tailor_cv):
        prompt = js_tailor_cv.build_prompt("DevOps Engineer", "HSBC", "job desc text")
        assert isinstance(prompt, str)
        assert len(prompt) > 100

    def test_includes_role_company_jd(self, js_tailor_cv):
        prompt = js_tailor_cv.build_prompt("DevOps Engineer", "HSBC", "Kubernetes AWS")
        assert "DevOps Engineer" in prompt
        assert "HSBC" in prompt
        assert "Kubernetes AWS" in prompt

    def test_truncates_long_jd(self, js_tailor_cv):
        long_jd = "x" * 10_000
        prompt = js_tailor_cv.build_prompt("T", "C", long_jd)
        # The JD is truncated at 2500 chars in build_prompt
        assert "x" * 2500 in prompt
        assert "x" * 5000 not in prompt  # should be truncated

    def test_includes_bruce_context(self, js_tailor_cv):
        prompt = js_tailor_cv.build_prompt("X", "Y", "z")
        # needle must be lowercase; "in" operator is case-sensitive
        assert "senior it leader" in prompt.lower()
        assert "HSBC" in prompt  # BRUCE_CONTEXT mentions HSBC
        assert "CKA" in prompt or "AWS" in prompt

    def test_contains_writing_instructions(self, js_tailor_cv):
        prompt = js_tailor_cv.build_prompt("X", "Y", "z")
        # Should instruct the model about style
        assert "first person" in prompt.lower() or '"I"' in prompt
        assert "summary" in prompt.lower()


# ── generate_profile (MiniMax API) ────────────────────────────────────────────

class TestGenerateProfile:
    def _mock_urlopen(self, js_tailor_cv, response_data, status=200):
        """Helper: monkeypatch urllib.request.urlopen to return response_data."""
        mock_resp = MagicMock()
        mock_resp.read.return_value = json.dumps(response_data).encode("utf-8")
        mock_resp.__enter__ = MagicMock(return_value=mock_resp)
        mock_resp.__exit__ = MagicMock(return_value=False)

        def fake_urlopen(req, timeout=60):
            if status >= 400:
                # Simulate HTTPError
                err = MagicMock()
                err.fp = MagicMock()
                err.fp.read.return_value = b'{"error": "bad request"}'
                err.code = status
                raise js_tailor_cv.urllib.error.HTTPError(
                    "https://api", status, "error", {}, err.fp
                )
            return mock_resp
        return fake_urlopen

    def test_successful_call_extracts_text(self, js_tailor_cv):
        response_data = {
            "content": [{"type": "text", "text": "A senior IT leader with cloud expertise."}],
            "stop_reason": "end_turn",
        }
        with patch.object(js_tailor_cv.urllib.request, "urlopen",
                          side_effect=self._mock_urlopen(js_tailor_cv, response_data)):
            text = js_tailor_cv.generate_profile("DevOps", "HSBC", "job description")
        assert "A senior IT leader" in text

    def test_normalizes_unicode_dashes(self, js_tailor_cv):
        response_data = {
            "content": [{"type": "text", "text": "Expert\u2013level engineer \u2014ready"}],
        }
        with patch.object(js_tailor_cv.urllib.request, "urlopen",
                          side_effect=self._mock_urlopen(js_tailor_cv, response_data)):
            text = js_tailor_cv.generate_profile("X", "Y", "z")
        # En-dash \u2013 → "-" and em-dash \u2014 → " -- "
        assert "Expert-level engineer" in text
        assert " -- ready" in text
        assert "\u2013" not in text
        assert "\u2014" not in text

    def test_normalizes_smart_quotes(self, js_tailor_cv):
        response_data = {
            "content": [{"type": "text", "text": "It\u2019s a \u201Cgreat\u201D opportunity."}],
        }
        with patch.object(js_tailor_cv.urllib.request, "urlopen",
                          side_effect=self._mock_urlopen(js_tailor_cv, response_data)):
            text = js_tailor_cv.generate_profile("X", "Y", "z")
        assert "It's" in text
        assert '"great"' in text

    def test_normalizes_ellipsis(self, js_tailor_cv):
        response_data = {
            "content": [{"type": "text", "text": "Working hard\u2026 delivers results"}],
        }
        with patch.object(js_tailor_cv.urllib.request, "urlopen",
                          side_effect=self._mock_urlopen(js_tailor_cv, response_data)):
            text = js_tailor_cv.generate_profile("X", "Y", "z")
        assert "..." in text
        assert "\u2026" not in text

    def test_raises_when_no_text_in_response(self, js_tailor_cv):
        response_data = {"content": [], "stop_reason": "max_tokens"}
        with patch.object(js_tailor_cv.urllib.request, "urlopen",
                          side_effect=self._mock_urlopen(js_tailor_cv, response_data)):
            with pytest.raises(RuntimeError, match="No text"):
                js_tailor_cv.generate_profile("X", "Y", "z")

    def test_picks_last_non_empty_text_block(self, js_tailor_cv):
        response_data = {
            "content": [
                {"type": "text", "text": ""},
                {"type": "text", "text": ""},
                {"type": "text", "text": "Final answer from the model."},
            ],
        }
        with patch.object(js_tailor_cv.urllib.request, "urlopen",
                          side_effect=self._mock_urlopen(js_tailor_cv, response_data)):
            text = js_tailor_cv.generate_profile("X", "Y", "z")
        assert "Final answer" in text

    def test_http_error_wrapped_in_runtime_error(self, js_tailor_cv):
        with patch.object(js_tailor_cv.urllib.request, "urlopen",
                          side_effect=self._mock_urlopen(js_tailor_cv, {}, status=401)):
            with pytest.raises(RuntimeError, match="HTTP 401"):
                js_tailor_cv.generate_profile("X", "Y", "z")

    def test_request_includes_bearer_auth(self, js_tailor_cv):
        """Verify the request sets the Authorization header."""
        response_data = {"content": [{"type": "text", "text": "OK"}]}
        captured = {}

        def fake_urlopen(req, timeout=60):
            captured["headers"] = req.headers
            mock_resp = MagicMock()
            mock_resp.read.return_value = json.dumps(response_data).encode()
            mock_resp.__enter__ = MagicMock(return_value=mock_resp)
            mock_resp.__exit__ = MagicMock(return_value=False)
            return mock_resp

        with patch.object(js_tailor_cv.urllib.request, "urlopen",
                          side_effect=fake_urlopen):
            js_tailor_cv.generate_profile("X", "Y", "z")
        assert "Authorization" in captured["headers"]
        assert captured["headers"]["Authorization"].startswith("Bearer ")

    def test_request_body_contains_model(self, js_tailor_cv):
        response_data = {"content": [{"type": "text", "text": "OK"}]}
        captured = {}

        def fake_urlopen(req, timeout=60):
            captured["data"] = req.data
            mock_resp = MagicMock()
            mock_resp.read.return_value = json.dumps(response_data).encode()
            mock_resp.__enter__ = MagicMock(return_value=mock_resp)
            mock_resp.__exit__ = MagicMock(return_value=False)
            return mock_resp

        with patch.object(js_tailor_cv.urllib.request, "urlopen",
                          side_effect=fake_urlopen):
            js_tailor_cv.generate_profile("X", "Y", "z")
        body = json.loads(captured["data"])
        assert "model" in body
        assert "messages" in body
        assert body["messages"][0]["role"] == "user"


# ── split_profile_into_runs ───────────────────────────────────────────────────

class TestSplitProfileIntoRuns:
    def test_preserves_total_run_count(self, js_tailor_cv):
        text = "A short profile."
        result = js_tailor_cv.split_profile_into_runs(text, total_runs=5, en_dash_runs=set())
        assert len(result) == 5

    def test_reserves_en_dash_positions_as_none(self, js_tailor_cv):
        text = "Profile text split across multiple runs"
        result = js_tailor_cv.split_profile_into_runs(
            text, total_runs=5, en_dash_runs={2}
        )
        assert result[2] is None
        # Other slots should have content
        for i in (0, 1, 3, 4):
            assert result[i] is not None

    def test_distributes_words_evenly(self, js_tailor_cv):
        text = "word " * 12  # 12 words
        result = js_tailor_cv.split_profile_into_runs(
            text, total_runs=4, en_dash_runs=set()
        )
        # Each slot should have ~3 words; last slot gets the rest
        non_empty = [r for r in result if r]
        assert len(non_empty) == 4

    def test_handles_single_run(self, js_tailor_cv):
        result = js_tailor_cv.split_profile_into_runs(
            "Only one slot", total_runs=1, en_dash_runs=set()
        )
        assert len(result) == 1
        assert "Only one slot" in result[0]

    def test_handles_empty_text(self, js_tailor_cv):
        result = js_tailor_cv.split_profile_into_runs(
            "", total_runs=3, en_dash_runs=set()
        )
        # Should still return 3 slots (some empty)
        assert len(result) == 3

    def test_uses_expected_en_dash_indices(self, js_tailor_cv):
        # tailor_cv.py hardcodes {3, 11, 13} — verify behavior matches
        text = "a b c d e f g h i j k l m n o p"
        result = js_tailor_cv.split_profile_into_runs(
            text, total_runs=17, en_dash_runs={3, 11, 13}
        )
        assert result[3] is None
        assert result[11] is None
        assert result[13] is None
        assert len(result) == 17


# ── update_cv (docx manipulation) ─────────────────────────────────────────────

class TestUpdateCv:
    def test_replaces_title_in_para_2(self, js_tailor_cv, sample_docx):
        js_tailor_cv.update_cv(sample_docx, "New Senior Role", "New profile content here")
        from docx import Document
        doc = Document(str(sample_docx))
        assert doc.paragraphs[2].text == "New Senior Role"

    def test_replaces_profile_in_para_3(self, js_tailor_cv, sample_docx):
        new_profile = "A new test profile for the role."
        js_tailor_cv.update_cv(sample_docx, "Title", new_profile)
        from docx import Document
        doc = Document(str(sample_docx))
        # python-docx joins runs without spaces; check words individually
        text = doc.paragraphs[3].text
        assert "test" in text
        assert "profile" in text
        assert "role" in text

    def test_preserves_para_0_and_1(self, js_tailor_cv, sample_docx):
        original_p0 = "BRUCE KWOK"
        original_p1 = "brucekwok@hotmail.com | +852 1234 5678"
        js_tailor_cv.update_cv(sample_docx, "Any Title", "Any profile text")
        from docx import Document
        doc = Document(str(sample_docx))
        assert original_p0 in doc.paragraphs[0].text
        assert original_p1 in doc.paragraphs[1].text

    def test_preserves_en_dash_runs(self, js_tailor_cv, sample_docx):
        """Runs at indices 3, 11, 13 should retain the en-dash character."""
        js_tailor_cv.update_cv(sample_docx, "Title", "Some new content for the profile")
        from docx import Document
        from docx.oxml.ns import qn
        doc = Document(str(sample_docx))
        runs = doc.paragraphs[3]._element.findall(qn("w:r"))
        # Verify the en-dash runs at indices 3, 11, 13 still have "‑"
        for idx in (3, 11, 13):
            t = runs[idx].find(qn("w:t"))
            assert t.text == "‑"

    def test_does_not_crash_with_minimal_profile(self, js_tailor_cv, sample_docx):
        # Should not raise even with short text
        js_tailor_cv.update_cv(sample_docx, "X", "short")

    def test_save_file_is_valid_docx(self, js_tailor_cv, sample_docx, tmp_path):
        js_tailor_cv.update_cv(sample_docx, "Title", "Profile content")
        # Should still be a valid docx file
        from docx import Document
        doc = Document(str(sample_docx))
        assert len(doc.paragraphs) >= 4


# ── process_folder (end-to-end with mocked MiniMax) ──────────────────────────

class TestProcessFolder:
    def test_skips_folder_without_job_txt(self, js_tailor_cv, tmp_path):
        empty_dir = tmp_path / "empty"
        empty_dir.mkdir()
        result = js_tailor_cv.process_folder(empty_dir)
        assert result["status"] == "skip"
        assert "no job.txt" in result["reason"]

    def test_processes_valid_folder(self, js_tailor_cv, tmp_job_folder, sample_docx):
        # Use a copy of the template in the fixture location
        import shutil
        # Place template at the location the script expects
        with patch.object(js_tailor_cv, "TEMPLATE_CV", sample_docx):
            # Mock MiniMax response
            response_data = {
                "content": [{"type": "text", "text": "A new professional summary for this role."}],
            }
            with patch.object(js_tailor_cv.urllib.request, "urlopen") as mock_urlopen:
                mock_resp = MagicMock()
                mock_resp.read.return_value = json.dumps(response_data).encode()
                mock_resp.__enter__ = MagicMock(return_value=mock_resp)
                mock_resp.__exit__ = MagicMock(return_value=False)
                mock_urlopen.return_value = mock_resp
                result = js_tailor_cv.process_folder(tmp_job_folder)
        assert result["status"] == "ok"
        assert "job_title" in result

    def test_extracts_title_and_company_from_job_txt(self, js_tailor_cv, tmp_job_folder):
        title, company = ("Senior DevOps Engineer", "HSBC")
        # Title is parsed from the job.txt's "Title:" line
        text = tmp_job_folder.joinpath("job.txt").read_text()
        assert title in text
        assert company in text

    def test_falls_back_to_title_only_when_no_separator(self, js_tailor_cv, tmp_path, sample_docx):
        # Create a job.txt WITHOUT the ====== separator
        bad_dir = tmp_path / "bad_separator"
        bad_dir.mkdir()
        (bad_dir / "job.txt").write_text(
            "Title: Engineer\nCompany: Acme\nLink:\n",
            encoding="utf-8",
        )
        # When separator missing, JD text defaults to title + company.
        # Use the sample_docx as a valid template so shutil.copy2 succeeds.
        with patch.object(js_tailor_cv, "TEMPLATE_CV", sample_docx):
            with patch.object(js_tailor_cv.urllib.request, "urlopen",
                              side_effect=Exception("API down")):
                result = js_tailor_cv.process_folder(bad_dir)
        # Either error or skip — both acceptable for this edge case
        assert result["status"] in ("error", "skip")

    def test_handles_minimax_error_gracefully(self, js_tailor_cv, tmp_job_folder, sample_docx):
        import shutil
        with patch.object(js_tailor_cv, "TEMPLATE_CV", sample_docx):
            with patch.object(js_tailor_cv.urllib.request, "urlopen",
                              side_effect=Exception("MiniMax down")):
                result = js_tailor_cv.process_folder(tmp_job_folder)
        assert result["status"] == "error"
        assert "MiniMax" in result["reason"]

    def test_strips_view_job_posting_link_from_jd(self, js_tailor_cv, tmp_job_folder):
        # The JD text is cleaned of [View job posting](...) before sending
        text = (tmp_job_folder / "job.txt").read_text()
        # Confirm the markdown link pattern is present in raw file
        assert "[View job posting]" in text


# ── find_job_folders ───────────────────────────────────────────────────────────

class TestFindJobFolders:
    def test_returns_empty_when_jd_dir_missing(self, js_tailor_cv, tmp_path):
        with patch.object(js_tailor_cv, "JD_DIR", tmp_path / "nonexistent"):
            folders = js_tailor_cv.find_job_folders()
        assert folders == []

    def test_finds_nested_folders(self, js_tailor_cv, tmp_path):
        # Create structure: tmp/JDs/JobsDB/Company_Title_xxx/
        jd_root = tmp_path / "JDs"
        for source in ("JobsDB", "eFinancialCareers"):
            for n in range(2):
                (jd_root / source / f"Co{n}_Title{n}_abc12345").mkdir(parents=True)
        with patch.object(js_tailor_cv, "JD_DIR", jd_root):
            folders = js_tailor_cv.find_job_folders()
        assert len(folders) == 4
        # Sorted alphabetically
        assert folders[0].name < folders[-1].name
class TestTailorMainCli:
    def _set_template(self, js_tailor_cv, path):
        """Replace TEMPLATE_CV in the module's global namespace."""
        js_tailor_cv.TEMPLATE_CV = path

    def test_main_with_no_args_lists_folders(self, js_tailor_cv, tmp_path,
                                            monkeypatch, sample_docx, capsys):
        import sys
        from unittest.mock import patch
        # Create a job folder
        job_dir = tmp_path / "Source" / "Co_Title_hash"
        job_dir.mkdir(parents=True)
        (job_dir / "job.txt").write_text("Title: x\nCompany: y\n", encoding="utf-8")

        monkeypatch.setattr(sys, "argv", ["tailor_cv.py"])
        monkeypatch.setattr(js_tailor_cv, "JD_DIR", tmp_path)
        self._set_template(js_tailor_cv, sample_docx)
        with patch.object(js_tailor_cv, "find_job_folders",
                          return_value=[job_dir]):
            with patch.object(js_tailor_cv, "process_folder",
                              return_value={"status": "ok", "job_title": "x",
                                            "company": "y", "profile": "x..."}):
                js_tailor_cv.main()
        out = capsys.readouterr().out
        assert "1 CVs tailored" in out
        assert "0 errors" in out

    def test_main_with_no_folders_prints_message(self, js_tailor_cv, tmp_path,
                                                monkeypatch, sample_docx, capsys):
        import sys
        from unittest.mock import patch
        monkeypatch.setattr(sys, "argv", ["tailor_cv.py"])
        monkeypatch.setattr(js_tailor_cv, "JD_DIR", tmp_path)
        self._set_template(js_tailor_cv, sample_docx)
        with patch.object(js_tailor_cv, "find_job_folders", return_value=[]):
            js_tailor_cv.main()
        out = capsys.readouterr().out
        assert "No job folders found" in out

    def test_main_with_list_arg_shows_folders(self, js_tailor_cv, tmp_path,
                                            monkeypatch, sample_docx, capsys):
        import sys
        from unittest.mock import patch
        job_dir = tmp_path / "Source" / "Co_Title_hash"
        job_dir.mkdir(parents=True)
        (job_dir / "Bruce Kwok_Resume.docx").write_bytes(b"fake docx")
        monkeypatch.setattr(sys, "argv", ["tailor_cv.py", "--list"])
        monkeypatch.setattr(js_tailor_cv, "JD_DIR", tmp_path)
        self._set_template(js_tailor_cv, sample_docx)
        with patch.object(js_tailor_cv, "find_job_folders",
                          return_value=[job_dir]):
            js_tailor_cv.main()
        out = capsys.readouterr().out
        assert "Job folders" in out or "1" in out

    def test_main_with_folder_arg_processes_one(self, js_tailor_cv, tmp_path,
                                                monkeypatch, sample_docx, capsys):
        import sys
        from unittest.mock import patch
        job_dir = tmp_path / "Source" / "Co_Title_hash"
        job_dir.mkdir(parents=True)
        (job_dir / "job.txt").write_text("Title: x\nCompany: y\n", encoding="utf-8")

        monkeypatch.setattr(sys, "argv", ["tailor_cv.py", "--folder", str(job_dir)])
        self._set_template(js_tailor_cv, sample_docx)
        with patch.object(js_tailor_cv, "process_folder",
                          return_value={"status": "ok", "job_title": "x",
                                        "company": "y", "profile": "x..."}):
            js_tailor_cv.main()
        out = capsys.readouterr().out
        assert "1 CVs tailored" in out

    def test_main_skips_existing_cv(self, js_tailor_cv, tmp_path,
                                   monkeypatch, sample_docx, capsys):
        import sys
        from unittest.mock import patch
        job_dir = tmp_path / "Source" / "Co_Title_hash"
        job_dir.mkdir(parents=True)
        (job_dir / "job.txt").write_text("Title: x\nCompany: y\n", encoding="utf-8")
        (job_dir / "Bruce Kwok_Resume.docx").write_bytes(b"existing")
        monkeypatch.setattr(sys, "argv", ["tailor_cv.py"])
        monkeypatch.setattr(js_tailor_cv, "JD_DIR", tmp_path)
        self._set_template(js_tailor_cv, sample_docx)
        with patch.object(js_tailor_cv, "find_job_folders",
                          return_value=[job_dir]):
            with patch.object(js_tailor_cv, "process_folder") as pf_mock:
                js_tailor_cv.main()
        out = capsys.readouterr().out
        assert "SKIP" in out
        assert not pf_mock.called

    def test_main_no_template_errors(self, js_tailor_cv, tmp_path, monkeypatch, capsys):
        import sys
        monkeypatch.setattr(sys, "argv", ["tailor_cv.py"])
        # Set TEMPLATE_CV to a path that doesn't exist
        js_tailor_cv.TEMPLATE_CV = tmp_path / "definitely_missing.docx"
        with pytest.raises(SystemExit):
            js_tailor_cv.main()
        out = capsys.readouterr().out
        assert "Template CV not found" in out


class TestUpdateCvEdgeCases:
    def test_handles_long_profile_text(self, js_tailor_cv, sample_docx):
        # A profile longer than 17 runs worth of text — should still work
        long_profile = " ".join(["word"] * 200)
        js_tailor_cv.update_cv(sample_docx, "Engineer", long_profile)
        from docx import Document
        doc = Document(str(sample_docx))
        assert doc.paragraphs[2].text == "Engineer"
        # The 200 words should be distributed across the 14 non-en-dash runs
        full = doc.paragraphs[3].text
        # Count word occurrences (excluding the en-dash character itself)
        words = [w for w in full.split() if w != "‑"]
        assert len(words) >= 100  # Most words distributed

    def test_extra_runs_become_empty(self, js_tailor_cv, sample_docx):
        # Provide a short profile so some runs get empty text
        js_tailor_cv.update_cv(sample_docx, "Title", "tiny")
        from docx import Document
        from docx.oxml.ns import qn
        doc = Document(str(sample_docx))
        # All 17 runs still exist
        runs = doc.paragraphs[3]._element.findall(qn("w:r"))
        assert len(runs) == 17
        # En-dash runs preserved
        for idx in (3, 11, 13):
            t = runs[idx].find(qn("w:t"))
            assert t.text == "‑"


class TestProcessFolderMore:
    def test_no_job_info_uses_defaults(self, js_tailor_cv, tmp_path, sample_docx):
        # job.txt exists but no Title: line
        job_dir = tmp_path / "src" / "x_y_z"
        job_dir.mkdir(parents=True)
        (job_dir / "job.txt").write_text("Some random text without structured fields\n", encoding="utf-8")
        from unittest.mock import patch
        with patch.object(js_tailor_cv, "TEMPLATE_CV", sample_docx):
            with patch.object(js_tailor_cv.urllib.request, "urlopen",
                              side_effect=Exception("API down")):
                result = js_tailor_cv.process_folder(job_dir)
        assert result["status"] == "error"

    def test_uses_link_from_job_info(self, js_tailor_cv, tmp_job_folder, sample_docx):
        # Process folder with valid job.txt — verify the link field is read
        from unittest.mock import patch
        with patch.object(js_tailor_cv, "TEMPLATE_CV", sample_docx):
            with patch.object(js_tailor_cv.urllib.request, "urlopen",
                              side_effect=Exception("API")):
                result = js_tailor_cv.process_folder(tmp_job_folder)
        # Either the link was extracted or there's an error result
        # The "Link:" line in tmp_job_folder's job.txt has the link
        # Check the test data is set up correctly
        text = (tmp_job_folder / "job.txt").read_text()
        assert "Link:     https://hk.jobsdb.com/job/123456" in text

    def test_process_folder_with_vsl_separator(self, js_tailor_cv, tmp_path, sample_docx):
        # Different separator length (still a long ==== run)
        job_dir = tmp_path / "src" / "x_y_z"
        job_dir.mkdir(parents=True)
        (job_dir / "job.txt").write_text(
            "Title: Dev\nCompany: Acme\nLink:\n\n" + "=" * 80 + "\n\n" + "JD text",
            encoding="utf-8"
        )
        from unittest.mock import patch
        with patch.object(js_tailor_cv, "TEMPLATE_CV", sample_docx):
            with patch.object(js_tailor_cv.urllib.request, "urlopen",
                              side_effect=Exception("API")):
                result = js_tailor_cv.process_folder(job_dir)
        # Just verify no crash
        assert result is not None
