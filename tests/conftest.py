"""
Pytest configuration & shared fixtures for tjl_live_futu (HK scanner).

Loads the scanner module via importlib (since it's a script with side-effect
imports that read futu modules and bind them as ft.*). Provides:

  - tjl_mod: the loaded tjl_live_futu module
  - make_quote / make_bars: deterministic synthetic data builders
  - fake_ctx: a Mock for futu.OpenQuoteContext with controllable behaviour
"""
from __future__ import annotations

import importlib.util
import os
import sys
import types
from unittest.mock import MagicMock

import numpy as np
import pandas as pd
import pytest


# ── Stub the futu package so we can import tjl_live_futu without OpenD ────────
@pytest.fixture(scope="session", autouse=True)
def stub_futu():
    """Inject a fake `futu` module + `futu.quote.open_quote_context` before
    importing tjl_live_futu. This prevents real network calls during tests."""

    futu_pkg = types.ModuleType("futu")
    futu_pkg.Market = types.SimpleNamespace(HK="HK")
    futu_pkg.SecurityType = types.SimpleNamespace(STOCK="STOCK")

    class _KLType:
        K_DAY = "K_DAY"
        K_30M = "K_30M"
        K_15M = "K_15M"
        K_5M = "K_5M"
        K_1M = "K_1M"

    class _SubType:
        QUOTE = "QUOTE"
        ORDER_BOOK = "ORDER_BOOK"

    class _OpenQuoteContext:
        def __init__(self, *a, **kw): pass
        def close(self): pass
        def request_history_kline(self, *a, **kw): return (0, None, None)
        def subscribe(self, *a, **kw): pass
        def get_stock_quote(self, *a, **kw): return (0, None, None)
        def get_market_snapshot(self, *a, **kw): return (0, [])

    futu_pkg.KLType = _KLType
    futu_pkg.SubType = _SubType
    futu_pkg.OpenQuoteContext = _OpenQuoteContext

    # futu.quote.open_quote_context sub-module path used by the scanner
    quote_mod = types.ModuleType("futu.quote")
    oqc_mod = types.ModuleType("futu.quote.open_quote_context")
    oqc_mod.OpenQuoteContext = _OpenQuoteContext
    oqc_mod.KLType = _KLType
    oqc_mod.SubType = _SubType

    sys.modules["futu"] = futu_pkg
    sys.modules["futu.quote"] = quote_mod
    sys.modules["futu.quote.open_quote_context"] = oqc_mod

    # Re-bind on the package object so `from futu.quote.open_quote_context import ...`
    # works on subsequent imports.
    futu_pkg.OpenQuoteContext = _OpenQuoteContext
    futu_pkg.KLType = _KLType
    futu_pkg.SubType = _SubType

    yield futu_pkg


# ── Load the scanner under test ───────────────────────────────────────────────
@pytest.fixture(scope="session")
def tjl_mod(stub_futu):
    """Import tjl_live_futu.py via importlib (it is a script, not a module)."""
    script = os.path.join(
        os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
        "tjl_live_futu.py",
    )
    spec = importlib.util.spec_from_file_location("tjl_live_futu", script)
    mod = importlib.util.module_from_spec(spec)
    sys.modules["tjl_live_futu"] = mod
    spec.loader.exec_module(mod)
    return mod


# ── Synthetic data builders ───────────────────────────────────────────────────
@pytest.fixture
def uptrend_bars():
    """80 bars trending up; EMAs will stack bullish (E9>E20>E50)."""
    closes = np.array([100.0 + 0.05 * i for i in range(80)])
    highs = closes + 0.30
    lows = closes - 0.30
    volumes = np.full(80, 1_000_000, dtype=int)
    return highs, lows, closes, volumes


@pytest.fixture
def downtrend_bars():
    closes = np.array([200.0 - 0.05 * i for i in range(80)])
    highs = closes + 0.30
    lows = closes - 0.30
    volumes = np.full(80, 1_000_000, dtype=int)
    return highs, lows, closes, volumes


@pytest.fixture
def flat_bars():
    closes = np.array([100.0] * 80)
    highs = closes + 0.50
    lows = closes - 0.50
    volumes = np.full(80, 1_000_000, dtype=int)
    return highs, lows, closes, volumes


def make_quote(code, price=100.0, prev_close=99.0, high_today=100.5,
               low_today=99.5, open_today=99.8, volume=1_000_000):
    """Return a single-row pandas DataFrame shaped like futu's quote DataFrame."""
    return pd.DataFrame([{
        "code": code,
        "last_price": price,
        "prev_close_price": prev_close,
        "high_price": high_today,
        "low_price": low_today,
        "open_price": open_today,
        "volume": volume,
    }])


@pytest.fixture
def fake_ctx():
    """MagicMock standing in for futu.OpenQuoteContext with sensible defaults."""
    ctx = MagicMock()
    ctx.subscribe.return_value = (0, None)
    ctx.close.return_value = (0, None)
    return ctx


# ════════════════════════════════════════════════════════════════════════════════
# job-search skill fixtures
# ════════════════════════════════════════════════════════════════════════════════
JOB_SKILL_DIR = "/Users/jaydensmac/.hermes/skills/job-search"
JOB_SCRIPT_DIR = f"{JOB_SKILL_DIR}/scripts"


def _js_load_module(name, file, extra_mock=False):
    """Load a job-search script via importlib."""
    if extra_mock:
        playwright_mod = types.ModuleType("playwright")
        sync_api = types.ModuleType("playwright.sync_api")
        sync_api.sync_playwright = MagicMock()
        playwright_mod.sync_api = sync_api
        sys.modules.setdefault("playwright", playwright_mod)
        sys.modules.setdefault("playwright.sync_api", sync_api)
    spec = importlib.util.spec_from_file_location(name, file)
    mod = importlib.util.module_from_spec(spec)
    sys.modules[name] = mod
    spec.loader.exec_module(mod)
    return mod


@pytest.fixture(scope="session")
def js_config():
    return _js_load_module("jobsearch_config", f"{JOB_SCRIPT_DIR}/config.py")


@pytest.fixture(scope="session")
def js_save_jds():
    return _js_load_module("jobsearch_save_jds", f"{JOB_SCRIPT_DIR}/save_jds.py")


@pytest.fixture(scope="session")
def js_build_profile():
    return _js_load_module("jobsearch_build_profile",
                           f"{JOB_SCRIPT_DIR}/build_profile.py")


@pytest.fixture(scope="session")
def js_scrape_all():
    return _js_load_module("jobsearch_scrape_all",
                           f"{JOB_SCRIPT_DIR}/scrape_all.py",
                           extra_mock=True)


@pytest.fixture(scope="session")
def js_tailor_cv():
    return _js_load_module("jobsearch_tailor_cv",
                           f"{JOB_SCRIPT_DIR}/tailor_cv.py")


@pytest.fixture(scope="session")
def js_update_tracker():
    return _js_load_module("jobsearch_update_tracker",
                           f"{JOB_SCRIPT_DIR}/update_tracker.py")


@pytest.fixture
def tmp_job_csv(tmp_path):
    """Build a small but realistic CSV of scraped jobs."""
    rows = [
        {"title": "Senior DevOps Engineer",
         "company": "HSBC",
         "location": "Hong Kong",
         "posted": "2d ago",
         "salary": "HK$80K-120K/mo",
         "source": "JobsDB",
         "link": "https://hk.jobsdb.com/job/123456"},
        {"title": "Junior Cloud Engineer",
         "company": "StartupX",
         "location": "Kwun Tong",
         "posted": "1d ago",
         "salary": "HK$30K-40K/mo",
         "source": "JobsDB",
         "link": "https://hk.jobsdb.com/job/789012"},
        {"title": "Lead SRE — AWS Kubernetes",
         "company": "eFinancial",
         "location": "Central",
         "posted": "3d ago",
         "salary": "HK$100K-150K/mo",
         "source": "eFinancialCareers",
         "link": "https://www.efinancialcareers.hk/job/345678"},
    ]
    csv_path = tmp_path / "jobs_combined.csv"
    import csv as _csv
    with csv_path.open("w", newline="", encoding="utf-8") as f:
        writer = _csv.DictWriter(f, fieldnames=list(rows[0].keys()))
        writer.writeheader()
        writer.writerows(rows)
    return csv_path


@pytest.fixture
def tmp_job_folder(tmp_path):
    """Build a job folder with a job.txt containing a typical JD layout."""
    src = "JobsDB"
    title = "Senior DevOps Engineer"
    company = "HSBC"
    link = "https://hk.jobsdb.com/job/123456"
    job_dir = tmp_path / src / f"{company}_{title}_abcd1234"
    job_dir.mkdir(parents=True)
    job_txt = job_dir / "job.txt"
    job_txt.write_text(
        f"Source:   {src}\n"
        f"Company:  {company}\n"
        f"Title:    {title}\n"
        f"Location: Hong Kong\n"
        f"Salary:   HK$80K-120K/mo\n"
        f"Posted:   2d ago\n"
        f"Link:     {link}\n"
        "\n"
        "=" * 60 + "\n"
        "\n"
        f"# {title} @ {company}\n"
        f"[View job posting]({link})\n"
        "\n"
        "## Full Job Description\n"
        "(auto-fetched — used for CV profile tailoring)\n"
        "\n"
        "We are hiring a Senior DevOps Engineer with strong AWS, Kubernetes, and "
        "Terraform experience to join our platform team.\n"
        "\n"
        "## Notes\n"
        "(Add your notes after reviewing the JD)\n",
        encoding="utf-8",
    )
    return job_dir


@pytest.fixture
def sample_docx(tmp_path):
    """Create a minimal valid docx that mimics Bruce_Kwok_Resume.docx structure.

    tailor_cv.py reads paragraphs[2] (title) and paragraphs[3] (profile with
    17 runs, en-dashes at indices 3, 11, 13).
    """
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.add_paragraph("BRUCE KWOK")
    doc.add_paragraph("brucekwok@hotmail.com | +852 1234 5678")
    doc.add_paragraph("Cloud & Platform Senior Lead")

    p3 = doc.add_paragraph()
    chunks = [
        "A senior IT leader with 15+ years of ",
        "experience in cloud infrastructure. ",
        "Certified CKA and AWS SA. ",
        "‑",  # en-dash run
        "Built and led high-performing ",
        "SRE teams. ",
        "Delivered audit-ready solutions. ",
        "Hands-on with Terraform, ",
        "Kubernetes, and CI/CD pipelines. ",
        "Background in financial services. ",
        "PMP and PRINCE2 certified. ",
        "‑",  # en-dash run
        "Strong communicator. ",
        "‑",  # en-dash run
        "Vendor management experience. ",
        "Multi-cloud governance. ",
        "Strategic thinker.",
    ]
    assert len(chunks) == 17, f"need 17 chunks, got {len(chunks)}"
    for txt in chunks:
        run = p3.add_run(txt)
        run.font.size = Pt(11)

    out = tmp_path / "Bruce Kwok_Resume.docx"
    doc.save(str(out))
    return out